from __future__ import annotations

from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "reports"
OUT.mkdir(parents=True, exist_ok=True)

INSIGHT_PATH = OUT / "supabase_exposure_insight_report_2026-08-15.docx"
TEAM_PATH = OUT / "team_report_section_4.2_supabase_exposure.docx"

BLUE = "16324F"
MID_BLUE = "2E74B5"
DARK_BLUE = "0B2545"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "5B6573"
RED = "9B1C1C"
PALE_RED = "FDECEC"
GOLD = "7A5A00"
PALE_GOLD = "FFF6D8"
GREEN = "176B45"
WHITE = "FFFFFF"
BLACK = "172033"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D6DCE5", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_fixed_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    set_run_font(run, 9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MID_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r_pr.extend([r_fonts, color, underline])
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_custom_numbering(doc, marker="•"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if marker == "•" else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker if marker == "•" else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id):
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    run = paragraph.add_run(text)
    set_run_font(run, 11, color=BLACK)
    return paragraph


def base_document(running_title):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, MID_BLUE, 16, 8),
        ("Heading 2", 13, MID_BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = running_title
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(3)
    for run in hp.runs:
        set_run_font(run, 8.5, bold=True, color=MUTED)
    p_pr = hp._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D6DCE5")
    border.append(bottom)
    p_pr.append(border)
    add_page_number(section.footer.paragraphs[0])
    return doc


def add_title_block(doc, eyebrow, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(eyebrow.upper())
    set_run_font(r, 9, bold=True, color=MID_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, 25, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    set_run_font(r, 12, color=MUTED)


def add_callout(doc, label, text, fill=PALE_BLUE, color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_fixed_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size="0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "  ")
    set_run_font(r, 10, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, 11, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_kpi_table(doc, items):
    table = doc.add_table(rows=2, cols=len(items))
    set_fixed_table_geometry(table, [9360 // len(items)] * len(items))
    set_table_borders(table, color="D6DCE5", size="6")
    for i, (value, label) in enumerate(items):
        set_cell_shading(table.cell(0, i), BLUE)
        p = table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        set_run_font(r, 20, bold=True, color=WHITE)
        set_cell_shading(table.cell(1, i), LIGHT_GRAY)
        p = table.cell(1, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, 9.5, bold=True, color=DARK_BLUE)
    return table


def add_table(doc, headers, rows, widths, header_fill=BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_fixed_table_geometry(table, widths)
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, 9, bold=True, color=WHITE)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if row_index % 2 == 1:
                set_cell_shading(cells[i], "F8FAFC")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, 9.2, color=BLACK)
        set_fixed_table_geometry(table, widths)
    return table


def add_source_line(doc, prefix, link_text=None, url=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(prefix)
    set_run_font(r, 8.5, color=MUTED)
    if link_text and url:
        add_hyperlink(p, link_text, url)
    return p


def build_insight_report():
    doc = base_document("PUBLIC WEB EXPOSURE RESEARCH · 2026-08-15")
    bullets = add_custom_numbering(doc)
    add_title_block(
        doc,
        "Supabase exposure insight",
        "공개 웹앱의 Supabase 고권한 자격증명 노출 조사",
        "공개 HTML·JavaScript 정적 분석 결과 | 2026년 8월 13–15일 (KST)",
    )
    add_callout(
        doc,
        "한 줄 결론",
        "426개 공개 URL 중 3개 도메인이 HIGH 검토 후보로 분류됐다. 이 가운데 2건은 실제 Supabase 프로젝트 URL과 고권한 키 형식이 같은 공개 JavaScript 번들에 함께 존재해 즉시 비공개 확인과 키 교체가 필요한 강한 정황이다.",
        PALE_RED,
        RED,
    )
    add_kpi_table(doc, [("426", "고유 URL 점검"), ("384", "정상 분석"), ("8", "수동 검토 필요"), ("3", "HIGH 후보")])
    add_source_line(doc, "분모: batch_summary.csv의 고유 URL 426개. 정상 분석 384개(90.1%), 오류 35개, 건너뜀 7개.")

    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph()
    r = p.add_run("판정: 부분 확인 · 신뢰도 중간")
    set_run_font(r, 12, bold=True, color=RED)
    p.add_run(" — 공개 번들에 고권한 자격증명 형태의 문자열이 존재한다는 사실은 확인됐다. 다만 키의 현재 유효성, 데이터베이스 접근 성공 여부, RLS 상태, 실제 개인정보 열람 여부는 시험하지 않았으므로 ‘개인정보 유출 확정’으로 표현해서는 안 된다.")
    for text in (
        "가장 강한 사례 2건: 공개 JS에서 Supabase 프로젝트 URL과 함께 `sb_secret_…` 후보 또는 `service_role` 역할 JWT 후보가 탐지됐다.",
        "추가 HIGH 1건: 자체 Supabase 호환/실습 성격으로 보이는 페이지에서 `JWT_SECRET` 및 `service_role` 문자열이 관찰됐으나 운영형 클라우드 프로젝트와 연결되는지는 미확인이다.",
        "`PRIVATE KEY` 탐지 1건은 설정 입력창의 예시문구로 보여 독립적인 유출 근거에서 제외했다.",
        "보고서와 공유본에는 원문 키와 대상 도메인을 싣지 않았다. 정확한 URL과 마스킹 문맥은 내부 엑셀의 ‘검토 후보’·‘상세 탐지’ 시트에서만 확인한다.",
    ):
        add_list_item(doc, text, bullets)

    doc.add_heading("1. 조사 범위와 방법", level=1)
    doc.add_paragraph(
        "Brave 검색 결과와 수동 입력 URL을 합쳐 공개 웹페이지를 수집한 뒤, 각 페이지의 HTML과 연결된 JavaScript 파일에서 Supabase 프로젝트 URL, 공개용 키, 고권한 키, 토큰 및 개인정보 형식 후보를 정규식과 문맥 점수로 분류했다. 접근은 일반 방문자가 내려받을 수 있는 공개 자산에 한정했다."
    )
    add_table(
        doc,
        ["단계", "수행 내용", "한계"],
        [
            ("발견", "128개 검색어, API 요청 806회, 발견 URL 408개", "검색엔진 색인과 검색어 편향 존재"),
            ("수집", "고유 URL 426개, HTML·연결 JS 점검", "로그인 뒤 화면·동적 API 응답 제외"),
            ("탐지", "형식·라벨·주변 문맥으로 신뢰도 산정", "형식 일치만으로 진짜 비밀값 확정 불가"),
            ("검토", "HIGH/MEDIUM/LOW 우선순위 부여", "키 유효성·DB 접근·RLS는 시험하지 않음"),
        ],
        [1300, 4300, 3760],
    )
    add_source_line(doc, "내부 근거: supabase_scan_report.xlsx, report_candidates.csv, results.csv (스냅샷 2026-08-15 12:56 KST).")

    doc.add_heading("2. 노출 퍼널", level=1)
    add_table(
        doc,
        ["구간", "도메인 수", "정상 분석 대비", "해석"],
        [
            ("정상 분석", "384", "100.0%", "오류·건너뜀 제외"),
            ("Supabase 신호 또는 민감 형식 후보", "115", "29.9%", "공개용 키·일반 연락처 포함 가능"),
            ("REVIEW_REQUIRED", "8", "2.1%", "분석가 확인이 필요한 문맥"),
            ("HIGH 검토 후보", "3", "0.78%", "고권한 자격증명 유형 후보 포함"),
            ("강한 정황", "2", "0.52%", "프로젝트 URL과 고권한 키가 같은 공개 번들에 공존"),
        ],
        [2600, 1500, 1700, 3560],
    )
    add_callout(
        doc,
        "해석 주의",
        "115개를 ‘유출 사이트’로 세면 안 된다. Supabase 프로젝트 URL과 publishable/anon 키는 브라우저에 포함될 수 있도록 설계된 공개 식별자다. 위험 판단은 키 종류와 RLS·권한 구성을 함께 봐야 한다.",
        PALE_GOLD,
        GOLD,
    )

    doc.add_heading("3. HIGH 후보 사례 분석", level=1)
    add_table(
        doc,
        ["사례", "공개 자산에서 관찰", "현재 판단", "미확인 사항"],
        [
            ("A · Render", "프로젝트 URL + `sb_secret_…` 후보가 동일 JS 번들에 존재", "강한 고권한 키 노출 정황", "유효성·DB 접근·피해"),
            ("B · Vercel", "프로젝트 URL + `service_role` 역할 JWT 후보가 동일 JS 번들에 존재", "강한 고권한 키 노출 정황", "유효성·DB 접근·피해"),
            ("C · Cloudflare Pages", "`JWT_SECRET` 라벨과 `service_role` 역할 JWT 후보가 공개 JS에 존재", "기술적 비밀값 노출 후보; 데모/로컬 구현 가능성", "운영 프로젝트 연결·유효성"),
        ],
        [1250, 3650, 2550, 1910],
    )
    add_source_line(doc, "대상 이름과 값은 비공개 처리했다. 정확한 행은 내부 엑셀에서 review_priority=HIGH로 필터링한다.")

    doc.add_heading("3.1 오탐을 분리한 이유", level=2)
    doc.add_paragraph(
        "한 HIGH 도메인에서 `PRIVATE KEY` 헤더가 함께 탐지됐지만, 마스킹 문맥상 사용자가 키를 붙여 넣는 입력창의 placeholder 예시였다. 따라서 이 항목은 ‘실제 개인키 노출’ 근거로 사용하지 않았다. 반면 `sb_secret_…` 및 역할이 `service_role`로 해석된 JWT는 Supabase 프로젝트 URL과 코드상 클라이언트 초기화 주변에 위치해 상대적으로 강한 정황으로 보았다."
    )

    doc.add_heading("4. 왜 위험한가", level=1)
    doc.add_paragraph(
        "Supabase 공식 문서에 따르면 publishable/anon 키는 낮은 권한의 공개용 키이며 RLS 정책의 보호를 받는다. 반대로 secret/service_role 키는 서버 전용 고권한 자격증명이다. 특히 service_role은 RLS를 우회하므로 프런트엔드에 포함되면 권한 경계 자체가 무너질 수 있다."
    )
    add_source_line(doc, "공식 기준: ", "Supabase — Understanding API keys", "https://supabase.com/docs/guides/getting-started/api-keys")
    add_table(
        doc,
        ["키/신호", "브라우저 공개", "의미"],
        [
            ("publishable / anon", "허용", "공개 식별자. 단, 모든 테이블의 RLS와 정책 검토가 필요"),
            ("secret / service_role", "금지", "서버 전용 고권한. 데이터 접근 권한 및 RLS 우회 위험"),
            ("프로젝트 URL", "대체로 허용", "단독으로는 비밀이 아니며 키 종류와 함께 판단"),
        ],
        [2300, 1700, 5360],
    )

    doc.add_heading("5. 설정 노출과 실제 피해 범위", level=1)
    add_table(
        doc,
        ["확인 단계", "이번 조사 상태", "보고서 표현"],
        [
            ("공개 번들에 문자열 존재", "확인", "고권한 자격증명 ‘노출 후보’ 관찰"),
            ("키가 현재 유효함", "미확인", "유효 키라고 단정하지 않음"),
            ("DB/Storage 접근 가능", "미시험", "RLS 미설정이라고 단정하지 않음"),
            ("개인정보 레코드 열람", "미확인", "개인정보 유출 건수로 집계하지 않음"),
            ("악용·피해 발생", "미확인", "침해사고로 단정하지 않음"),
        ],
        [2700, 1900, 4760],
    )

    doc.add_heading("6. 권고 조치", level=1)
    for text in (
        "즉시: 운영 주체에게 비공개 채널로 알리고, 노출 의심 secret/service_role 키를 교체·폐기한다.",
        "즉시: 고권한 작업을 브라우저 코드에서 제거하고 서버 또는 Supabase Edge Function으로 이동한다.",
        "24시간 이내: 배포 시점 이후의 Auth·Database·Storage·Edge Function 로그를 검토해 비정상 접근 여부를 확인한다.",
        "24시간 이내: 모든 테이블의 RLS 활성화와 anon/authenticated 정책을 별도로 점검한다. service_role 노출 여부와 RLS 검증은 서로 대체되지 않는다.",
        "중장기: `VITE_`, `NEXT_PUBLIC_`, `PUBLIC_` 환경변수에 고권한 값이 들어가면 빌드를 실패시키는 검사를 CI에 추가한다.",
        "공개 보고: 키 원문·정확한 엔드포인트·개인정보 표본은 싣지 않고, 조치 완료 전에는 대상 이름을 최소화한다.",
    ):
        add_list_item(doc, text, bullets)

    doc.add_heading("7. 한계와 재현 기준", level=1)
    doc.add_paragraph(
        "이번 조사는 수동 침투시험이 아니라 공개 자산의 수동적 정적 분석이다. 따라서 결과는 특정 시점의 프런트엔드 배포 상태를 보여줄 뿐이며, 키 교체 여부나 백엔드 권한 상태는 반영하지 못할 수 있다. 동일 보고서를 재현하려면 동일한 URL 목록, 탐지 규칙 버전, 스캔 시각과 결과 해시를 보존해야 한다."
    )
    add_source_line(doc, "구조 참고: ", "SupaExplorer — State of Supabase Exposure Across Vibe-Coding Apps", "https://supaexplorer.com/cybersecurity-insight-report-january-2026")
    add_source_line(doc, "주의: 위 외부 보고서는 스키마 접근성까지 시험했지만, 이번 조사는 Supabase API·DB·Storage를 호출하지 않았으므로 수치와 판정은 직접 비교할 수 없다.")
    doc.add_heading("판정 기준", level=2)
    add_table(
        doc,
        ["판정", "적용 기준"],
        [
            ("확인", "공개 자산에 해당 문자열이 존재하고, 독립적인 문맥이 유형 판정을 지지"),
            ("부분 확인", "공개 노출은 확인됐으나 유효성·권한·피해 범위 일부가 미확인"),
            ("미확인", "형식 일치 외의 근거가 부족하거나 대안 설명을 배제할 수 없음"),
            ("오탐 제외", "placeholder·문서 예시·압축 코드 숫자열 등 비밀값이 아닌 문맥 확인"),
        ],
        [2100, 7260],
    )
    doc.add_heading("증거 취급", level=2)
    for text in (
        "원문 자격증명은 재현·보고서에 복사하지 않고 마스킹 값과 evidence_hash만 보존한다.",
        "대상 URL은 내부 엑셀에 제한하고 외부 공유본에서는 사례 코드와 호스팅 유형만 사용한다.",
        "후속 확인 시 스캔 시각, 공개 자산 URL, 탐지 규칙 버전, 조치 여부를 함께 기록한다.",
    ):
        add_list_item(doc, text, bullets)

    doc.save(INSIGHT_PATH)


def build_team_report():
    doc = base_document("팀 보고서 · 클라우드·웹 백엔드 설정 오류")
    bullets = add_custom_numbering(doc)
    add_title_block(
        doc,
        "Team report section",
        "4.2 클라우드·웹 백엔드 설정 오류",
        "Supabase 공개 프런트엔드 자격증명 노출 조사 요약 | 2026-08-15",
    )
    add_callout(
        doc,
        "보고용 결론",
        "공개 웹페이지 426개를 조사한 결과 HIGH 검토 후보 3개를 찾았다. 이 중 2건은 공개 JavaScript에 Supabase 프로젝트 주소와 서버 전용 고권한 키 형식이 함께 포함된 강한 정황이다. 다만 실제 데이터 접근이나 개인정보 유출은 확인하지 않았다.",
        PALE_RED,
        RED,
    )

    doc.add_heading("4.2.1 공개 노출 구조", level=1)
    doc.add_paragraph(
        "Supabase와 Firebase 같은 BaaS(Backend as a Service)는 로그인, 데이터베이스, 파일 저장소 같은 서버 기능을 웹 개발자가 빠르게 붙일 수 있게 해주는 서비스다. 이번 조사는 그중 Supabase를 대상으로 했다."
    )
    doc.add_paragraph(
        "웹 브라우저는 화면을 실행하기 위해 HTML과 JavaScript 파일을 내려받는다. 따라서 그 파일 안에 넣은 값은 방문자가 볼 수 있다. Supabase의 프로젝트 URL과 publishable/anon 키는 원래 공개 사용을 전제로 하지만, secret/service_role 키는 서버에만 있어야 한다. 이 구분을 놓치면 다음 구조로 노출이 발생한다."
    )
    add_table(
        doc,
        ["정상 구조", "오류 구조"],
        [
            ("브라우저 → 공개용 키 → RLS 정책 → 허용된 데이터", "브라우저 → 서버용 고권한 키 → RLS 우회 가능"),
            ("중요한 작업은 서버/Edge Function에서 처리", "고권한 키를 `VITE_`·`PUBLIC_` 변수로 배포"),
        ],
        [4680, 4680],
    )
    add_source_line(doc, "기준: ", "Supabase 공식 API 키 설명", "https://supabase.com/docs/guides/getting-started/api-keys")

    doc.add_heading("4.2.2 노출 유형 분석", level=1)
    add_kpi_table(doc, [("426", "점검 URL"), ("384", "정상 분석"), ("8", "수동 검토"), ("3", "HIGH 후보")])
    doc.add_paragraph(
        "HIGH는 ‘피해 확정’이 아니라 분석가가 먼저 확인해야 하는 우선순위다. 3개 사례를 다시 살펴본 결과는 다음과 같다."
    )
    add_table(
        doc,
        ["사례", "발견 내용", "판단"],
        [
            ("A", "프로젝트 URL + `sb_secret_…` 후보가 공개 JS에 함께 존재", "강한 고권한 키 노출 정황"),
            ("B", "프로젝트 URL + `service_role` JWT 후보가 공개 JS에 함께 존재", "강한 고권한 키 노출 정황"),
            ("C", "`JWT_SECRET` 및 `service_role` 문자열이 공개 JS에 존재", "노출 후보이나 데모/로컬 구현 가능성 확인 필요"),
        ],
        [1000, 4960, 3400],
    )
    doc.add_paragraph(
        "또한 `PRIVATE KEY` 탐지 1건은 입력창에 표시된 예시문구로 보여 실제 개인키 노출에서는 제외했다. 정확한 사이트와 마스킹 증거는 supabase_scan_report.xlsx에서 review_priority 열을 HIGH로 필터링하면 확인할 수 있다."
    )

    doc.add_heading("4.2.3 설정 노출과 실제 피해 범위 구분", level=1)
    doc.add_paragraph(
        "이번 조사에서 확인한 것은 ‘공개 파일에 고권한 자격증명으로 보이는 값이 들어 있다’는 사실이다. 키를 사용해 데이터베이스에 접속하거나 데이터를 조회하지 않았으므로, RLS 미설정이나 실제 개인정보 유출까지 확인한 것은 아니다."
    )
    add_table(
        doc,
        ["구분", "결과"],
        [
            ("공개 설정/자격증명 노출", "3개 HIGH 후보, 그중 2개 강한 정황"),
            ("키 유효성", "확인하지 않음"),
            ("RLS 미설정·우회 성공", "확인하지 않음"),
            ("개인정보 레코드 노출", "확인하지 않음"),
            ("실제 피해", "확인되지 않음"),
        ],
        [3300, 6060],
    )
    add_callout(
        doc,
        "권장 보고 문장",
        "‘공개 웹 자산에서 Supabase 고권한 자격증명 노출 후보 3개 도메인을 탐지했으며, 2건은 프로젝트 URL과 고권한 키가 동일 번들에 존재하는 강한 정황으로 확인됐다. 실제 데이터 접근 및 개인정보 피해는 검증 범위에 포함하지 않았다.’ 조치: 운영 주체에 비공개 통지 후 키 교체, 서버 측 이전, 로그·RLS 점검을 권고한다.",
        PALE_BLUE,
        DARK_BLUE,
    )
    doc.save(TEAM_PATH)


def audit(path):
    doc = Document(path)
    section = doc.sections[0]
    assert round(section.page_width / 914400, 3) == 8.5
    assert round(section.page_height / 914400, 3) == 11.0
    assert all(round(v / 914400, 3) == 1.0 for v in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin))
    normal = doc.styles["Normal"].paragraph_format
    assert normal.space_after.pt == 6
    assert abs(normal.line_spacing - 1.10) < 0.001
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and tbl_w.get(qn("w:w")) == "9360"
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tbl_ind is not None and tbl_ind.get(qn("w:w")) == "120"
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "***MASKED***" not in all_text
    assert "eyJh" not in all_text
    assert "supabase-lite.pages.dev" not in all_text
    assert "tracking-clients.onrender.com" not in all_text
    assert "vue-chat-lovat.vercel.app" not in all_text


if __name__ == "__main__":
    build_insight_report()
    build_team_report()
    audit(INSIGHT_PATH)
    audit(TEAM_PATH)
    print(INSIGHT_PATH)
    print(TEAM_PATH)
